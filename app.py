import os
import re
import json
import datetime

import streamlit as st
import pdfplumber
from docx import Document
from odf.opendocument import load
from odf.text import P
from odf.teletype import extractText

import gspread
from google.oauth2.service_account import Credentials

# Nouveau SDK Gemini
from google import genai
from google.genai import types

# ==========================================
# 0. CONFIG GEMINI (NUAGE)
# ==========================================

if "GEMINI_API_KEY" in st.secrets:
    client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])
else:
    api_key_env = os.getenv("GEMINI_API_KEY")
    if not api_key_env:
        st.error("GEMINI_API_KEY manquante (Secrets Streamlit ou variable d'environnement).")
    client = genai.Client(api_key=api_key_env)

# ==========================================
# 1. CONSTANTES (CANON CORRIGÉ)
# ==========================================

CONSTANTES_SAGA = {
    "Jonas": """ÂGE: ~17 ans. Prénom complet: Jonas. Surnom: Grizzly (donné par Léo).
Orientation: Gay. Couple exclusif avec Léo (Moineau).
Aura: VERTE (phosphorescent, braise froide dans les crises).
Manifestation interne: "Colère" — créature verte, yeux de charbon, fumée noire, surgit de sa poitrine lors des rages extrêmes.
Capacités: chasseur expert (lance, couteau, fusil, pièges), tir méthodique et précis, instinct de survie.
Traits physiques: yeux verts, cheveux bruns, cicatrice à la gorge, épaules étroites, clavicule visible.
Talisman: petite boîte de métal.
Thèmes: culpabilité du chasseur, protection des siens, Colère comme ombre intérieure.""",

    "Léo": """ÂGE: ~15-17 ans selon le tome. Prénom complet: Léo. Surnom: Moineau (donné par Jonas).
Orientation: Bisexuel. Couple exclusif avec Jonas (Grizzly).
Aura: BLEUE (halo bleu, liens de lumière bleue).
Capacités: voit les fils/liens lumineux entre les gens; la Louve de lumière blanche sort de sa poitrine (gardienne);
  yeux qui brillent lors des visions; peut "arrêter le temps" dans les moments de crise intense.
Traits physiques: cheveux blonds, cicatrice fine sur la lèvre, yeux lumineux lors des visions.
Talisman: la Louve de lumière (intérieure).
Thèmes: protection, voir ce que les autres ne voient pas, lien entre les membres du groupe.""",

    "Zack": """ÂGE: 15 ans (Tome 2). Prénom choisi: Zackary. Ancien surnom: Gaz. Prénom légal d'origine: Albert (révélé par Hélène, T3).
Orientation: en découverte active — pas d'étiquette fixe dans le texte; intimité avec Jonas, Léo, Jade, Autyss selon les tomes.
Aura: ROUGE CHAUD (braise, incandescent — jamais noir).
Capacités physiques: membranes/ailes (flancs + entre jambes = "queue d'hirondelle"), deuxième paupière translucide,
  vol/planeur, électrostatique involontaire (gaz corporel + étincelle = combustion). Les ailes peuvent être blessées/déchirées.
Traits physiques: yeux noirs avec reflets noisette-rouille (après transformation par la Louve), corps très mince,
  cicatrices d'attaches de membranes, sourire "qui n'appartient qu'à dehors".
Talismans (jamais quittés): louveteau de pierre (sculpté par Jonas, T2) + sac en tissu avec fleur brodée + breloque-couteau.
Mère: Hélène (retrouvée T3), le décrit comme "enfant phénix".
Thèmes: corps passager (pas outil), apprendre le consentement, trouver son nom propre, trauma de la Cathédrale.""",

    "Jade": """Rôle: cheffe/matriarche de terrain, figure de stabilité du groupe.
Orientation: lien fort et progressif avec Zack; relation physique désirée mais à son rythme.
Aura: JAUNE/DORÉE (ligne jaune dans les visions de Léo).
Arme principale: fronde (elle porte aussi celle de Zack quand il en a besoin).
Traits: posture droite même sous pression, sens tactique, protège les plus vulnérables.
Thèmes: ne pas s'effacer pour le groupe, apprendre à dire non et oui avec clarté.""",

    "Autyss": """Rôle: stratège, "chirurgien", pensée en colonnes mentales.
Aura: VIOLETTE (ligne violette/craie violette dans les visions de Léo).
Capacités: analyse tactique extrême, dissociation contrôlée, sutures/soins médicaux, raisonnement en tableaux.
Besoins: constantes sensorielles pour ne pas "exploser" (fleur brodée, pierre, métal, eau chaude).
Talisman emprunté: sac de Zack (prêté dans les moments de crise).
Traits: peu de contact physique voulu initialement, s'exprime en poèmes lors des moments importants, observe avant d'agir.
Thèmes: trouver des mots pour les émotions, colonnes vs chaos, apprendre à ressentir sans disséquer.""",

    "SAGA": """Thèmes centraux: Reconstruction post-trauma ("corps passager, jamais outil").
L'Ombre (force antagoniste), la Cathédrale (organisation ennemie), la Voix (réseau de résistance).
Règles du groupe (T3): base Jonas–Léo intouchable; si flou = non; jalousie nommée avant qu'elle pourrisse;
  pas de sacrifice en douce; pas de sexe en secret; rappeler la règle avant de punir.
Antagonistes: Cité de pierre, Luceduort (commandos volants), Cathédrale.
Lieux clés: Badlands (désert rouge, tornades), Village Soleil, Forge, relais caravane.""",
}

# ==========================================
# 2. GOOGLE SHEETS (optionnel)
# ==========================================

SHEET_ID = "189e8EDBteW2bk-6XQMqz5CbDN7g2_CC-VY238jnC98I"

def connecter_et_obtenir_onglet(nom_onglet):
    try:
        if "GCP_JSON_BRUT" not in st.secrets:
            return None
        json_info = json.loads(st.secrets["GCP_JSON_BRUT"], strict=False)
        scope = [
            "https://spreadsheets.google.com/feeds",
            "https://www.googleapis.com/auth/drive",
        ]
        creds = Credentials.from_service_account_info(json_info, scopes=scope)
        client_gs = gspread.authorize(creds)
        ss = client_gs.open_by_key(SHEET_ID)
        try:
            return ss.worksheet(nom_onglet)
        except gspread.exceptions.WorksheetNotFound:
            nouvel_onglet = ss.add_worksheet(title=nom_onglet, rows="2000", cols="5")
            nouvel_onglet.append_row(["Date", "Diagnostic", "Type"])
            return nouvel_onglet
    except Exception as e:
        st.error(f"Erreur Sheets : {e}")
        return None

# ==========================================
# 3. IA – APPEL GEMINI (3 Flash)
# ==========================================

def appel_ia(prompt: str, temperature: float = 0.1) -> str:
    """
    Appel unique à Gemini 3 Flash via google-genai.
    """
    try:
        response = client.models.generate_content(
            model="gemini-2.0-flash",  # à remplacer par le nom exact de Gemini 3 Flash quand dispo
            contents=prompt,
            config=types.GenerateContentConfig(
                temperature=temperature,
                max_output_tokens=4096,
            ),
        )
        return response.text or ""
    except Exception as e:
        return f"❌ Erreur Gemini : {e}"

# ==========================================
# 4. FICHIERS – EXTRACTION & CHAPITRES
# ==========================================

def extraire_texte(f) -> str:
    """Lit PDF / DOCX / ODT et renvoie le texte brut."""
    try:
        if f.name.endswith(".pdf"):
            with pdfplumber.open(f) as pdf:
                return "
".join([p.extract_text() or "" for p in pdf.pages])
        elif f.name.endswith(".docx"):
            doc = Document(f)
            return "
".join([p.text for p in doc.paragraphs])
        elif f.name.endswith(".odt"):
            odt_doc = load(f)
            return "
".join([extractText(p) for p in odt_doc.getElementsByType(P)])
        return ""
    except Exception:
        return ""

def decouper_chapitres(texte: str):
    pattern = r'(?i)chapitres+d+'
    segments = re.split(pattern, texte)
    titres = re.findall(pattern, texte)
    chaps = []

    if len(segments) > 0 and len(segments[0].strip()) > 100:
        chaps.append({"titre": "Prologue", "contenu": segments[0].strip()})

    for i, t in enumerate(titres):
        if i + 1 < len(segments):
            chaps.append({"titre": t, "contenu": segments[i + 1].strip()})
    return chaps

# ==========================================
# 5. INTERFACE STREAMLIT (NUAGE)
# ==========================================

st.set_page_config(page_title="Archiviste Stable (Gemini SDK)", layout="wide")
st.title("🛡️ L'Archiviste V17.x (Roc de Pierre, Gemini 3 Flash)")

if "db" not in st.session_state:
    st.session_state.db = {"chapitres": [], "ready": False}

f_inputs = st.sidebar.file_uploader(
    "Charger les fichiers (tomes de la Saga)",
    type=["pdf", "docx", "odt"],
    accept_multiple_files=True,
)

if f_inputs and not st.session_state.db["ready"]:
    if st.button("🚀 Scanner la Saga"):
        full_txt = ""
        for f in f_inputs:
            full_txt += extraire_texte(f) + "

"
        full_txt = full_txt.replace("’", "'").replace("œ", "oe")
        st.session_state.db["chapitres"] = decouper_chapitres(full_txt)
        st.session_state.db["ready"] = True
        st.rerun()

if st.session_state.db["ready"]:
    perso = st.sidebar.selectbox(
        "Personnage",
        options=list(CONSTANTES_SAGA.keys()),
        help="Sélectionne un personnage ou 'SAGA' pour une vue globale."
    )

    indices = (
        [
            i for i, c in enumerate(st.session_state.db["chapitres"])
            if perso.lower() in c["contenu"].lower()
        ]
        if perso != "SAGA"
        else list(range(len(st.session_state.db["chapitres"])))
    )

    sel = st.multiselect(
        "Chapitres à analyser",
        options=indices,
        format_func=lambda x: st.session_state.db["chapitres"][x]["titre"],
    )

    if sel and st.button("🧠 Analyser"):
        txt_focus = "
".join(
            [st.session_state.db["chapitres"][i]["contenu"][:10000] for i in sel]
        )
        canon = CONSTANTES_SAGA[perso]
        prompt = f"""
[DIAGNOSTIC ARCHIVISTE STABLE]

PERSONNAGE : {perso}
CANON (à respecter strictement, cohérent avec les 4 tomes) :

{canon}

CONSIGNES :
- Analyse le personnage dans ces extraits (somatique, réactions, souveraineté, liens, trajectoire).
- Ne contredit jamais le canon ci-dessus.
- Si une information n'est pas présente dans les extraits, écris exactement : "Non mentionné dans le manuscrit.".
- Si le manuscrit contredit le canon, signale-le calmement, en détaillant : "INCOHÉRENCE DÉTECTÉE" + citation courte + page/chapitre si possible.
- Utilise des sous-titres clairs (Physique, Aura, Psyché, Liens, Trauma, Évolution, Incohérences éventuelles).

EXTRAITS SÉLECTIONNÉS :
{txt_focus}
"""
        with st.spinner("L'IA consulte les archives..."):
            resultat = appel_ia(prompt, temperature=0.1)
            st.markdown(resultat)
            st.session_state["dernier_resultat"] = resultat
            st.session_state["dernier_perso"] = perso

    if "dernier_resultat" in st.session_state and st.button("💾 Archiver dans Sheets"):
        onglet = connecter_et_obtenir_onglet(
            st.session_state.get("dernier_perso", "SAGA")
        )
        if onglet:
            date_now = datetime.datetime.now().strftime("%d/%m/%Y %H:%M")
            onglet.append_row(
                [date_now, st.session_state["dernier_resultat"], "ARCHIVISTE STABLE V17.x"]
            )
            st.success("✅ Diagnostic archivé dans Google Sheets.")

if st.sidebar.button("🗑️ Reset"):
    st.session_state.db = {"chapitres": [], "ready": False}
    st.rerun()    "Jade": """Rôle: cheffe/matriarche de terrain, figure de stabilité du groupe.
Orientation: lien fort et progressif avec Zack; relation physique désirée mais à son rythme.
Aura: JAUNE/DORÉE (ligne jaune dans les visions de Léo).
Arme principale: fronde (elle porte aussi celle de Zack quand il en a besoin).
Traits: posture droite même sous pression, sens tactique, protège les plus vulnérables.
Thèmes: ne pas s'effacer pour le groupe, apprendre à dire non et oui avec clarté.""",

    "Autyss": """Rôle: stratège, "chirurgien", pensée en colonnes mentales.
Aura: VIOLETTE (ligne violette/craie violette dans les visions de Léo).
Capacités: analyse tactique extrême, dissociation contrôlée, sutures/soins médicaux, raisonnement en tableaux.
Besoins: constantes sensorielles pour ne pas "exploser" (fleur brodée, pierre, métal, eau chaude).
Talisman emprunté: sac de Zack (prêté dans les moments de crise).
Traits: peu de contact physique voulu initialement, s'exprime en poèmes lors des moments importants, observe avant d'agir.
Thèmes: trouver des mots pour les émotions, colonnes vs chaos, apprendre à ressentir sans disséquer.""",

    "SAGA": """Thèmes centraux: Reconstruction post-trauma; souveraineté du corps ("corps passager, jamais outil");
  l'Ombre (force antagoniste), la Cathédrale (organisation ennemie), la Voix (réseau de résistance);
  règles de consentement inventées par le groupe autour du feu.
Règles du groupe (T3): base Jonas–Léo intouchable; si flou = non; jalousie nommée avant qu'elle pourrisse;
  pas de sacrifice en douce; pas de sexe en secret; rappeler la règle avant de punir.
Antagonistes: Cité de pierre, Luceduort (commandos volants), Cathédrale.
Lieux clés: Badlands (désert rouge, tornades), Village Soleil, Forge, relais caravane.""",
}

# ==========================================
# 2. GOOGLE SHEETS (optionnel)
# ==========================================

SHEET_ID = "189e8EDBteW2bk-6XQMqz5CbDN7g2_CC-VY238jnC98I"

def connecter_et_obtenir_onglet(nom_onglet):
    try:
        if "GCP_JSON_BRUT" not in st.secrets:
            return None
        json_info = json.loads(st.secrets["GCP_JSON_BRUT"], strict=False)
        scope = [
            "https://spreadsheets.google.com/feeds",
            "https://www.googleapis.com/auth/drive",
        ]
        creds = Credentials.from_service_account_info(json_info, scopes=scope)
        client_gs = gspread.authorize(creds)
        ss = client_gs.open_by_key(SHEET_ID)
        try:
            return ss.worksheet(nom_onglet)
        except gspread.exceptions.WorksheetNotFound:
            nouvel_onglet = ss.add_worksheet(title=nom_onglet, rows="2000", cols="5")
            nouvel_onglet.append_row(["Date", "Diagnostic", "Type"])
            return nouvel_onglet
    except Exception as e:
        st.error(f"Erreur Sheets : {e}")
        return None

# ==========================================
# 3. IA – APPEL GEMINI (3 Flash)
# ==========================================

def appel_ia(prompt: str, temperature: float = 0.1) -> str:
    """
    Appel unique à Gemini 3 Flash via google-genai.
    """
    try:
        response = client.models.generate_content(
            model="gemini-2.0-flash",  # remplace ici par le nom exact de Gemini 3 Flash quand dispo
            contents=prompt,
            config=types.GenerateContentConfig(
                temperature=temperature,
                max_output_tokens=4096,
            ),
        )
        return response.text or ""
    except Exception as e:
        return f"❌ Erreur Gemini : {e}"

# ==========================================
# 4. FICHIERS – EXTRACTION & CHAPITRES
# ==========================================

def extraire_texte(f) -> str:
    """Lit PDF / DOCX / ODT et renvoie le texte brut."""
    try:
        if f.name.endswith(".pdf"):
            with pdfplumber.open(f) as pdf:
                return "
".join([p.extract_text() or "" for p in pdf.pages])
        elif f.name.endswith(".docx"):
            doc = Document(f)
            return "
".join([p.text for p in doc.paragraphs])
        elif f.name.endswith(".odt"):
            odt_doc = load(f)
            return "
".join([extractText(p) for p in odt_doc.getElementsByType(P)])
        return ""
    except Exception:
        return ""

def decouper_chapitres(texte: str):
    pattern = r'(?i)chapitres+d+'
    segments = re.split(pattern, texte)
    titres = re.findall(pattern, texte)
    chaps = []

    if len(segments) > 0 and len(segments[0].strip()) > 100:
        chaps.append({"titre": "Prologue", "contenu": segments[0].strip()})

    for i, t in enumerate(titres):
        if i + 1 < len(segments):
            chaps.append({"titre": t, "contenu": segments[i + 1].strip()})
    return chaps

# ==========================================
# 5. INTERFACE STREAMLIT (NUAGE)
# ==========================================

st.set_page_config(page_title="Archiviste Stable (Gemini SDK)", layout="wide")
st.title("🛡️ L'Archiviste V17.x (Roc de Pierre, Gemini 3 Flash)")

if "db" not in st.session_state:
    st.session_state.db = {"chapitres": [], "ready": False}

f_inputs = st.sidebar.file_uploader(
    "Charger les fichiers (tomes de la Saga)",
    type=["pdf", "docx", "odt"],
    accept_multiple_files=True,
)

if f_inputs and not st.session_state.db["ready"]:
    if st.button("🚀 Scanner la Saga"):
        full_txt = ""
        for f in f_inputs:
            full_txt += extraire_texte(f) + "

"
        full_txt = full_txt.replace("’", "'").replace("œ", "oe")
        st.session_state.db["chapitres"] = decouper_chapitres(full_txt)
        st.session_state.db["ready"] = True
        st.rerun()

if st.session_state.db["ready"]:
    perso = st.sidebar.selectbox(
        "Personnage",
        options=list(CONSTANTES_SAGA.keys()),
        help="Sélectionne un personnage ou 'SAGA' pour une vue globale."
    )

    indices = (
        [
            i for i, c in enumerate(st.session_state.db["chapitres"])
            if perso.lower() in c["contenu"].lower()
        ]
        if perso != "SAGA"
        else list(range(len(st.session_state.db["chapitres"])))
    )

    sel = st.multiselect(
        "Chapitres à analyser",
        options=indices,
        format_func=lambda x: st.session_state.db["chapitres"][x]["titre"],
    )

    if sel and st.button("🧠 Analyser"):
        txt_focus = "
".join(
            [st.session_state.db["chapitres"][i]["contenu"][:10000] for i in sel]
        )
        canon = CONSTANTES_SAGA[perso]
        prompt = f"""
[DIAGNOSTIC ARCHIVISTE STABLE]

PERSONNAGE : {perso}
CANON (à respecter strictement, cohérent avec les 4 tomes) :

{canon}

CONSIGNES :
- Analyse le personnage dans ces extraits (somatique, réactions, souveraineté, liens, trajectoire).
- Ne contredit jamais le canon ci-dessus.
- Si une information n'est pas présente dans les extraits, écris exactement : "Non mentionné dans le manuscrit.".
- Si le manuscrit contredit le canon, signale-le calmement, en détaillant : "INCOHÉRENCE DÉTECTÉE" + citation courte + page/chapitre si possible.
- Utilise des sous-titres clairs (Physique, Aura, Psyché, Liens, Trauma, Évolution, Incohérences éventuelles).

EXTRAITS SÉLECTIONNÉS :
{txt_focus}
"""
        with st.spinner("L'IA consulte les archives..."):
            resultat = appel_ia(prompt, temperature=0.1)
            st.markdown(resultat)
            st.session_state["dernier_resultat"] = resultat
            st.session_state["dernier_perso"] = perso

    # Archivage optionnel dans Sheets
    if "dernier_resultat" in st.session_state and st.button("💾 Archiver dans Sheets"):
        onglet = connecter_et_obtenir_onglet(
            st.session_state.get("dernier_perso", "SAGA")
        )
        if onglet:
            date_now = datetime.datetime.now().strftime("%d/%m/%Y %H:%M")
            onglet.append_row(
                [date_now, st.session_state["dernier_resultat"], "ARCHIVISTE STABLE V17.x"]
            )
            st.success("✅ Diagnostic archivé dans Google Sheets.")

if st.sidebar.button("🗑️ Reset"):
    st.session_state.db = {"chapitres": [], "ready": False}
    st.rerun()"
        "Talisman: la Louve de lumière (intérieure).
"
        "Thèmes: protection, voir ce que les autres ne voient pas, lien entre les membres du groupe."
    ),
    "Zack": (
        "ÂGE: 15 ans (Tome 2). Prénom choisi: Zackary. Ancien surnom: Gaz. Prénom légal d'origine: Albert (révélé par Hélène, T3).
"
        "Orientation: en découverte active — pas d'étiquette fixe dans le texte; intimité avec Jonas, Léo, Jade, Autyss selon les tomes.
"
        "Aura: ROUGE CHAUD (braise, incandescent — jamais noir).
"
        "Capacités physiques: membranes/ailes (flancs + entre jambes = 'queue d'hirondelle'), deuxième paupière translucide,
"
        "  vol/planeur, électrostatique involontaire (gaz corporel + étincelle = combustion). Les ailes peuvent être blessées/déchirées.
"
        "Traits physiques: yeux noirs avec reflets noisette-rouille (après transformation par la Louve), corps très mince,
"
        "  cicatrices d'attaches de membranes, sourire 'qui n'appartient qu'à dehors'.
"
        "Talismans (jamais quittés): louveteau de pierre (sculpté par Jonas, T2) + sac en tissu avec fleur brodée + breloque-couteau.
"
        "Mère: Hélène (retrouvée T3), le décrit comme 'enfant phénix'.
"
        "Thèmes: corps passager (pas outil), apprendre le consentement, trouver son nom propre, trauma de la Cathédrale."
    ),
    "Jade": (
        "Rôle: cheffe/matriarche de terrain, figure de stabilité du groupe.
"
        "Orientation: lien fort et progressif avec Zack; relation physique désirée mais à son rythme.
"
        "Aura: JAUNE/DORÉE (ligne jaune dans les visions de Léo).
"
        "Arme principale: fronde (elle porte aussi celle de Zack quand il en a besoin).
"
        "Traits: posture droite même sous pression, sens tactique, protège les plus vulnérables.
"
        "Thèmes: ne pas s'effacer pour le groupe, apprendre à dire non et oui avec clarté."
    ),
    "Autyss": (
        "Rôle: stratège, 'chirurgien', pensée en colonnes mentales.
"
        "Aura: VIOLETTE (ligne violette/craie violette dans les visions de Léo).
"
        "Capacités: analyse tactique extrême, dissociation contrôlée, sutures/soins médicaux, raisonnement en tableaux.
"
        "Besoins: constantes sensorielles pour ne pas 'exploser' (fleur brodée, pierre, métal, eau chaude).
"
        "Talisman emprunté: sac de Zack (prêté dans les moments de crise).
"
        "Traits: peu de contact physique voulu initialement, s'exprime en poèmes lors des moments importants, observe avant d'agir.
"
        "Thèmes: trouver des mots pour les émotions, colonnes vs chaos, apprendre à ressentir sans disséquer."
    ),
    "SAGA": (
        "Thèmes centraux: Reconstruction post-trauma; souveraineté du corps ('corps passager, jamais outil');
"
        "  l'Ombre (force antagoniste), la Cathédrale (organisation ennemie), la Voix (réseau de résistance);
"
        "  règles de consentement inventées par le groupe autour du feu.
"
        "Règles du groupe (T3): base Jonas–Léo intouchable; si flou = non; jalousie nommée avant qu'elle pourrisse;
"
        "  pas de sacrifice en douce; pas de sexe en secret; rappeler la règle avant de punir.
"
        "Antagonistes: Cité de pierre, Luceduort (commandos volants), Cathédrale.
"
        "Lieux clés: Badlands (désert rouge, tornades), Village Soleil, Forge, relais caravane."
    ),
}

# ==========================================
# 2. GOOGLE SHEETS (optionnel)
# ==========================================

SHEET_ID = "189e8EDBteW2bk-6XQMqz5CbDN7g2_CC-VY238jnC98I"

def connecter_et_obtenir_onglet(nom_onglet):
    try:
        if "GCP_JSON_BRUT" not in st.secrets:
            return None
        json_info = json.loads(st.secrets["GCP_JSON_BRUT"], strict=False)
        scope = [
            "https://spreadsheets.google.com/feeds",
            "https://www.googleapis.com/auth/drive",
        ]
        creds = Credentials.from_service_account_info(json_info, scopes=scope)
        client_gs = gspread.authorize(creds)
        ss = client_gs.open_by_key(SHEET_ID)
        try:
            return ss.worksheet(nom_onglet)
        except gspread.exceptions.WorksheetNotFound:
            nouvel_onglet = ss.add_worksheet(title=nom_onglet, rows="2000", cols="5")
            nouvel_onglet.append_row(["Date", "Diagnostic", "Type"])
            return nouvel_onglet
    except Exception as e:
        st.error(f"Erreur Sheets : {e}")
        return None

# ==========================================
# 3. IA – APPEL GEMINI (3 Flash)
# ==========================================

def appel_ia(prompt: str, temperature: float = 0.1) -> str:
    """
    Appel unique à Gemini 3 Flash via google-genai.
    """
    try:
        response = client.models.generate_content(
            model="gemini-2.0-flash",  # remplace ici par le nom exact de Gemini 3 Flash quand dispo
            contents=prompt,
            config=types.GenerateContentConfig(
                temperature=temperature,
                max_output_tokens=4096,
            ),
        )
        return response.text or ""
    except Exception as e:
        return f"❌ Erreur Gemini : {e}"

# ==========================================
# 4. FICHIERS – EXTRACTION & CHAPITRES
# ==========================================

def extraire_texte(f) -> str:
    """Lit PDF / DOCX / ODT et renvoie le texte brut."""
    try:
        if f.name.endswith(".pdf"):
            with pdfplumber.open(f) as pdf:
                return "
".join([p.extract_text() or "" for p in pdf.pages])
        elif f.name.endswith(".docx"):
            doc = Document(f)
            return "
".join([p.text for p in doc.paragraphs])
        elif f.name.endswith(".odt"):
            odt_doc = load(f)
            return "
".join([extractText(p) for p in odt_doc.getElementsByType(P)])
        return ""
    except Exception:
        return ""

def decouper_chapitres(texte: str):
    pattern = r'(?i)chapitres+d+'
    segments = re.split(pattern, texte)
    titres = re.findall(pattern, texte)
    chaps = []

    if len(segments) > 0 and len(segments[0].strip()) > 100:
        chaps.append({"titre": "Prologue", "contenu": segments[0].strip()})

    for i, t in enumerate(titres):
        if i + 1 < len(segments):
            chaps.append({"titre": t, "contenu": segments[i + 1].strip()})
    return chaps

# ==========================================
# 5. INTERFACE STREAMLIT (NUAGE)
# ==========================================

st.set_page_config(page_title="Archiviste Stable (Gemini SDK)", layout="wide")
st.title("🛡️ L'Archiviste V17.x (Roc de Pierre, Gemini 3 Flash)")

if "db" not in st.session_state:
    st.session_state.db = {"chapitres": [], "ready": False}

f_inputs = st.sidebar.file_uploader(
    "Charger les fichiers (tomes de la Saga)",
    type=["pdf", "docx", "odt"],
    accept_multiple_files=True,
)

if f_inputs and not st.session_state.db["ready"]:
    if st.button("🚀 Scanner la Saga"):
        full_txt = ""
        for f in f_inputs:
            full_txt += extraire_texte(f) + "

"
        full_txt = full_txt.replace("’", "'").replace("œ", "oe")
        st.session_state.db["chapitres"] = decouper_chapitres(full_txt)
        st.session_state.db["ready"] = True
        st.rerun()

if st.session_state.db["ready"]:
    perso = st.sidebar.selectbox(
        "Personnage",
        options=list(CONSTANTES_SAGA.keys()),
        help="Sélectionne un personnage ou 'SAGA' pour une vue globale."
    )

    indices = [
        i for i, c in enumerate(st.session_state.db["chapitres"])
        if perso.lower() in c['contenu'].lower()
    ] if perso != "SAGA" else list(range(len(st.session_state.db["chapitres"])))

    sel = st.multiselect(
        "Chapitres à analyser",
        options=indices,
        format_func=lambda x: st.session_state.db["chapitres"][x]['titre'],
    )

    if sel and st.button("🧠 Analyser"):
        txt_focus = "
".join(
            [st.session_state.db["chapitres"][i]['contenu'][:10000] for i in sel]
        )
        canon = CONSTANTES_SAGA[perso]
        prompt = f"""
[DIAGNOSTIC ARCHIVISTE STABLE]

PERSONNAGE : {perso}
CANON (à respecter strictement, cohérent avec les 4 tomes) :

{canon}

CONSIGNES :
- Analyse le personnage dans ces extraits (somatique, réactions, souveraineté, liens, trajectoire).
- Ne contredit jamais le canon ci-dessus.
- Si une information n'est pas présente dans les extraits, écris exactement : "Non mentionné dans le manuscrit.".
- Si le manuscrit contredit le canon, signale-le calmement, en détaillant : "INCOHÉRENCE DÉTECTÉE" + citation courte + page/chapitre si possible.
- Utilise des sous-titres clairs (Physique, Aura, Psyché, Liens, Trauma, Évolution, Incohérences éventuelles).

EXTRAITS SÉLECTIONNÉS :
{txt_focus}
"""
        with st.spinner("L'IA consulte les archives..."):
            resultat = appel_ia(prompt, temperature=0.1)
            st.markdown(resultat)
            st.session_state["dernier_resultat"] = resultat
            st.session_state["dernier_perso"] = perso

    # Archivage optionnel dans Sheets
    if "dernier_resultat" in st.session_state and st.button("💾 Archiver dans Sheets"):
        onglet = connecter_et_obtenir_onglet(st.session_state.get("dernier_perso", "SAGA"))
        if onglet:
            date_now = datetime.datetime.now().strftime("%d/%m/%Y %H:%M")
            onglet.append_row(
                [date_now, st.session_state["dernier_resultat"], "ARCHIVISTE STABLE V17.x"]
            )
            st.success("✅ Diagnostic archivé dans Google Sheets.")

if st.sidebar.button("🗑️ Reset"):
    st.session_state.db = {"chapitres": [], "ready": False}
    st.rerun()
